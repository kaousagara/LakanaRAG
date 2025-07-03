from __future__ import annotations
import json
import os
import re
import time
from typing import List, Tuple, Deque
from collections import deque
from dataclasses import asdict
from docx import Document

from .base import QueryParam

class ToTNode:
    __slots__ = ('question', 'answer', 'depth', 'children')
    def __init__(self, question: str, depth: int):
        self.question = question
        self.answer = None
        self.depth = depth
        self.children = []

async def _parse_json(text: str) -> list[str]:
    match = re.search(r"\[.*\]", text.strip(), re.DOTALL)
    json_text = match.group(0).strip() if match else text.strip()
    try:
        data = json.loads(json_text)
        if isinstance(data, list):
            return [str(q) for q in data]
    except Exception:
        pass
    return []

async def _determine_depth(query: str, rag) -> int:
    """Détermine dynamiquement la profondeur max basée sur la complexité de la requête"""
    prompt = (
        "Évalue la complexité de cette requête et détermine la profondeur de recherche nécessaire "
        "(1=simple, 2=modérée, 3=complexe, 4=très complexe). Réponds UNIQUEMENT par un entier.\n\n"
        f"Requête: {query}\n\n"
        "Justification : Les sujets complexes nécessitent plus de sous-questions et d'approfondissements."
    )
    
    response = await rag.aquery(
        prompt,
        QueryParam(mode="naive", response_type="Value", stream=False),
        system_prompt=(
            "En tant qu'expert en analyse sémantique, évalue la complexité des requêtes. "
            "Fournis uniquement un chiffre entre 1 et 4 sans aucun commentaire."
        )
    )
    
    try:
        depth = int(response.strip())
        return max(1, min(4, depth))
    except (ValueError, TypeError):
        return 2 if len(query.split()) > 10 else 1

async def _generate_subqueries(query: str, rag) -> list[str]:
    plan_prompt = (
        "Décompose la requête principale en 2-4 sous-questions précises et complètes qui nécessitent "
        "des réponses détaillées et structurées. Chaque sous-question doit couvrir un aspect distinct "
        "du sujet principal.\n\n"
        f"Requête principale : {query}\n\n"
        "Format de réponse UNIQUEMENT : [\"sous-question 1\", \"sous-question 2\"]"
    )
    plan_text = await rag.aquery(
        plan_prompt,
        QueryParam(mode="naive", response_type="JSON", stream=False),
        system_prompt=(
            "Expert en analyse thématique : décompose les sujets complexes en sous-questions pertinentes "
            "qui permettent des développements approfondis."
        )
    )
    return await _parse_json(plan_text) or [query]

async def _generate_followups(question: str, answer: str, rag) -> list[str]:
    prompt = (
        "Génère 2 questions de suivi approfondies basées sur la question et sa réponse complète. "
        "Les nouvelles questions doivent explorer des aspects complémentaires ou demander des précisions "
        "sur des points spécifiques pour développer davantage l'analyse.\n\n"
        "### Question originale :\n"
        f"{question}\n\n"
        "### Réponse complète :\n"
        f"{answer}\n\n"
        "Format de réponse UNIQUEMENT : [\"question-suite 1\", \"question-suite 2\"]"
    )
    follow = await rag.aquery(
        prompt,
        QueryParam(mode="naive", response_type="JSON", stream=False),
        system_prompt=(
            "Spécialiste en approfondissement thématique : crée des questions de suivi "
            "qui permettent d'étendre l'analyse de manière cohérente et détaillée."
        )
    )
    return await _parse_json(follow) or []

async def _evaluate_thought(thought: str, context: str, rag) -> float:
    prompt = (
        "Évalue la pertinence et le potentiel de développement de cette question "
        "pour générer une réponse complète et structurée. Score entre 0 (hors-sujet) "
        "et 1 (excellent potentiel de développement).\n\n"
        "Contexte d'analyse :\n"
        f"{context}\n\n"
        "Question à évaluer :\n"
        f"{thought}\n\n"
        "Critères :\n"
        "- Pertinence par rapport au contexte\n"
        "- Potentiel pour une réponse détaillée\n"
        "- Originalité de l'angle d'approche\n"
        "- Contribution à la compréhension globale\n"
        "Réponds UNIQUEMENT par le score numérique."
    )
    response = await rag.aquery(
        prompt,
        QueryParam(mode="naive", response_type="Value", stream=False),
        system_prompt=(
            "Évaluateur expert : analyse la qualité des questions de recherche. "
            "Fournis uniquement un score entre 0 et 1 sans aucun commentaire."
        ),
    )
    try:
        return max(0.0, min(1.0, float(response.strip())))
    except:
        return 0.0

async def _select_thoughts(thoughts: List[str], context: str, rag, top_k: int) -> List[str]:
    if len(thoughts) <= top_k:
        return thoughts
    
    evaluations = []
    for thought in thoughts:
        score = await _evaluate_thought(thought, context, rag)
        evaluations.append((thought, score))
    
    evaluations.sort(key=lambda x: x[1], reverse=True)
    return [thought for thought, _ in evaluations[:top_k]]

async def _answer_question(question: str, rag, param: QueryParam) -> str:
    sub_param = QueryParam(**asdict(param))
    sub_param.mode = "hybrid"
    sub_param.stream = False
    
    # Prompt pour réponse complète et structurée
    full_prompt = (
        f"{question}\n\n"
        "Instructions :\n"
        "- Fournis une réponse complète, détaillée et bien structurée\n"
        "- Développe chaque point de manière approfondie\n"
        "- Utilise des paragraphes organisés avec une progression logique\n"
        "- Inclus si nécessaire des éléments de contexte pertinents\n"
        "- Évite les réponses concises ou fragmentées\n"
        "- Privilégie la clarté et l'exhaustivité"
    )
    
    ans = await rag.aquery(
        full_prompt,
        sub_param,
        system_prompt=(
            "Expert en rédaction analytique : produit des réponses complètes, "
            "structurées et riches en informations. Développe chaque point de manière approfondie "
            "avec une progression logique et des explications détaillées."
        )
    )
    return ans if isinstance(ans, str) else str(ans)

def _create_docx(content: str, working_dir: str) -> str:
    doc = Document()
    for line in content.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
    
    report_dir = os.path.join(working_dir, "reports")
    os.makedirs(report_dir, exist_ok=True)
    file_path = os.path.join(report_dir, f"deepsearch_{int(time.time())}.docx")
    doc.save(file_path)
    return file_path

def _format_report(query: str, qa_pairs: List[Tuple[str, str]]) -> str:
    text_lines = [f"# Rapport d'analyse approfondie : {query}", ""]
    text_lines.extend(["## Synthèse complète des investigations", ""])
    
    for idx, (q, a) in enumerate(qa_pairs, 1):
        text_lines.append(f"### Question d'analyse {idx} : {q}")
        text_lines.append("")
        text_lines.append(a)
        text_lines.append("")
    
    text_lines.append("## Conclusion analytique intégrée")
    text_lines.append("Cette analyse approfondie a exploré les différentes dimensions du sujet "
                      "à travers une investigation détaillée et structurée. Les réponses "
                      "complètes fournissent une compréhension exhaustive du thème initial.")
    return "\n".join(text_lines)

async def deepsearch_query(query: str, rag, param: QueryParam) -> str:
    # Détermination dynamique de la profondeur
    MAX_DEPTH = await _determine_depth(query, rag)
    print(f"Profondeur déterminée: {MAX_DEPTH} pour la requête: {query[:50]}...")
    
    # Paramètres adaptatifs
    MAX_INITIAL = max(2, min(4, MAX_DEPTH))
    MAX_FOLLOW = max(1, min(3, MAX_DEPTH - 1))

    # Construction de l'arbre
    root = ToTNode(query, depth=0)
    queue: Deque[ToTNode] = deque([root])
    
    while queue:
        node = queue.popleft()
        
        if node.depth > 0:
            node.answer = await _answer_question(node.question, rag, param)
        
        if node.depth < MAX_DEPTH:
            if node.depth == 0:
                children_questions = await _generate_subqueries(node.question, rag)
                selected = await _select_thoughts(
                    children_questions, 
                    node.question, 
                    rag, 
                    min(MAX_INITIAL, len(children_questions))
                )
            else:
                children_questions = await _generate_followups(node.question, node.answer, rag)
                context = f"{node.question}\n\n{node.answer}"
                selected = await _select_thoughts(
                    children_questions,
                    context,
                    rag,
                    min(MAX_FOLLOW, len(children_questions))
                )
            
            for q in selected:
                child = ToTNode(q, depth=node.depth + 1)
                node.children.append(child)
                queue.append(child)

    # Collecte des résultats
    qa_pairs = []
    collect_queue: Deque[ToTNode] = deque([root])
    while collect_queue:
        node = collect_queue.popleft()
        if node.depth > 0:
            qa_pairs.append((node.question, node.answer or ""))
        for child in node.children:
            collect_queue.append(child)
    
    # Génération du rapport
    report_text = _format_report(query, qa_pairs)
    return _create_docx(report_text, rag.working_dir)
